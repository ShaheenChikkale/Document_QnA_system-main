"""
Document Processing Service
Handles document upload, text extraction, and preprocessing
"""
import os
import uuid
import aiofiles
from typing import Optional, Tuple, Dict, Any
from pathlib import Path
from PyPDF2 import PdfReader
from docx import Document as DocxDocument
from langdetect import detect, LangDetectException

from app.utils.chunking import create_chunker
from app.utils.ocr import create_ocr_processor
from app.config import get_settings


class DocumentProcessor:
    """Process and extract text from various document formats"""
    
    def __init__(self):
        self.settings = get_settings()
        self.chunker = create_chunker(
            chunk_size=self.settings.chunk_size,
            chunk_overlap=self.settings.chunk_overlap
        )
        self.ocr_processor = create_ocr_processor()
        self.upload_dir = Path("data/uploads")
        self.upload_dir.mkdir(parents=True, exist_ok=True)
    
    async def process_document(
        self,
        file_content: bytes,
        filename: str
    ) -> Tuple[str, Dict[str, Any], str]:
        """
        Process uploaded document
        
        Args:
            file_content: File content bytes
            filename: Original filename
            
        Returns:
            Tuple of (document_id, metadata, extracted_text)
        """
        # Generate unique document ID
        document_id = str(uuid.uuid4())
        file_extension = Path(filename).suffix.lower()
        
        # Save file temporarily
        temp_path = self.upload_dir / f"{document_id}{file_extension}"
        async with aiofiles.open(temp_path, 'wb') as f:
            await f.write(file_content)
        
        # Extract text based on file type
        extracted_text, page_count = await self._extract_text(
            str(temp_path),
            file_extension
        )
        
        # Detect language
        language = self._detect_language(extracted_text)
        
        # Create metadata
        metadata = {
            "document_id": document_id,
            "filename": filename,
            "file_size": len(file_content),
            "file_extension": file_extension,
            "page_count": page_count,
            "language": language,
            "file_path": str(temp_path)
        }
        
        return document_id, metadata, extracted_text
    
    async def _extract_text(
        self,
        file_path: str,
        file_extension: str
    ) -> Tuple[str, Optional[int]]:
        """
        Extract text from document
        
        Args:
            file_path: Path to file
            file_extension: File extension
            
        Returns:
            Tuple of (extracted_text, page_count)
        """
        if file_extension == '.pdf':
            return await self._extract_from_pdf(file_path)
        elif file_extension == '.docx':
            return await self._extract_from_docx(file_path)
        elif file_extension == '.txt':
            return await self._extract_from_txt(file_path)
        elif file_extension in ['.png', '.jpg', '.jpeg']:
            return await self._extract_from_image(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
    
    async def _extract_from_pdf(self, file_path: str) -> Tuple[str, int]:
        """Extract text from PDF"""
        try:
            reader = PdfReader(file_path)
            page_count = len(reader.pages)
            
            # Try normal extraction
            text_parts = []
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                if text.strip():
                    text_parts.append(f"--- Page {i + 1} ---\n{text}")
            
            extracted_text = "\n\n".join(text_parts)
            
            # If very little text, try OCR
            if len(extracted_text.strip()) < 100:
                print(f"PDF appears to be scanned, applying OCR...")
                extracted_text = self.ocr_processor.extract_text_from_pdf_with_ocr(file_path)
            
            return extracted_text, page_count
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {str(e)}")
    
    async def _extract_from_docx(self, file_path: str) -> Tuple[str, int]:
        """Extract text from DOCX"""
        try:
            doc = DocxDocument(file_path)
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            text = "\n\n".join(paragraphs)
            page_count = len(doc.paragraphs) // 20 or 1  # Rough estimate
            return text, page_count
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {str(e)}")
    
    async def _extract_from_txt(self, file_path: str) -> Tuple[str, None]:
        """Extract text from TXT"""
        try:
            async with aiofiles.open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                text = await f.read()
            return text, None
        except Exception as e:
            raise ValueError(f"Failed to read text file: {str(e)}")
    
    async def _extract_from_image(self, file_path: str) -> Tuple[str, int]:
        """Extract text from image using OCR"""
        try:
            text = self.ocr_processor.extract_text_from_image(file_path)
            return text, 1
        except Exception as e:
            raise ValueError(f"Failed to extract text from image: {str(e)}")
    
    def _detect_language(self, text: str) -> str:
        """Detect text language"""
        if not text or len(text.strip()) < 10:
            return "unknown"
        
        try:
            return detect(text)
        except LangDetectException:
            return "unknown"
    
    def create_chunks(
        self,
        text: str,
        metadata: Dict[str, Any]
    ) -> list:
        """
        Create chunks from text
        
        Args:
            text: Text to chunk
            metadata: Document metadata
            
        Returns:
            List of chunks with metadata
        """
        return self.chunker.chunk_text(text, metadata)
    
    async def delete_document(self, file_path: str):
        """Delete document file"""
        try:
            if os.path.exists(file_path):
                os.remove(file_path)
        except Exception as e:
            print(f"Failed to delete file {file_path}: {str(e)}")


# Singleton instance
_document_processor = None

def get_document_processor() -> DocumentProcessor:
    """Get document processor singleton"""
    global _document_processor
    if _document_processor is None:
        _document_processor = DocumentProcessor()
    return _document_processor
